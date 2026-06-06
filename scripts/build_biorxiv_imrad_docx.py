from __future__ import annotations

from pathlib import Path

import build_manuscript_docx as builder
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]


builder.SOURCE = (
    ROOT
    / "PUBLICATION"
    / "0-preprint-bioRxiv"
    / "manuscript"
    / "cytomove-preprint-biorxiv-imrad.md"
)
builder.OUT = (
    ROOT
    / "PUBLICATION"
    / "0-preprint-bioRxiv"
    / "manuscript"
    / "Cytomove_preprint_bioRxiv_IMRAD.docx"
)


def configure_preprint_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(12)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    for name, size, before, after in [
        ("Heading 1", 12, 12, 6),
        ("Heading 2", 12, 10, 4),
        ("Heading 3", 12, 8, 4),
    ]:
        style = doc.styles[name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def add_preprint_cover(doc):
    title = doc.add_paragraph()
    title.alignment = builder.WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run(builder.read_working_title())
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)

    doc.add_page_break()


builder.configure_styles = configure_preprint_styles
builder.add_cover = add_preprint_cover


original_build_docx = builder.build_docx


def build_docx_with_black_headings():
    original_build_docx()
    doc = builder.Document(str(builder.OUT))
    for paragraph in doc.paragraphs:
        if paragraph.style.name.startswith("Heading"):
            for run in paragraph.runs:
                run.font.name = "Times New Roman"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
                run.font.size = Pt(12)
                run.font.bold = True
                run.font.color.rgb = RGBColor(0, 0, 0)
    doc.save(builder.OUT)


builder.build_docx = build_docx_with_black_headings


if __name__ == "__main__":
    builder.build_docx()
