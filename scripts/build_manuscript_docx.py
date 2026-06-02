from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageOps


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "manuscript-cytomove-submission.md"
OUT = ROOT / "docs" / "Cytomove_manuscript_submission.docx"
FIGURES = ROOT / "docs" / "manuscript_figures"
OPTIMIZED_FIGURES = ROOT / "docs" / "manuscript_figures" / "_docx_optimized"


FIGURE_MAP = {
    "Figure 1": FIGURES / "figure_1_app_workspace.png",
    "Figure 2": FIGURES / "figure_2_representative_overlays.png",
    "Figure 3": FIGURES / "figure_3_area_agreement.png",
    "Figure 4": FIGURES / "figure_4_whad_timecourse.png",
    "Figure 5": FIGURES / "figure_5_visual_comparison.png",
    "Supplementary Figure S1": FIGURES / "figure_s1_validation_summary.png",
    "Supplementary Figure S2": FIGURES / "figure_s2_csma_11_frame_visual_audit.png",
}


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, width_dxa: int = 9360) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")


def configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def apply_inline_markup(paragraph, text: str) -> None:
    text = text.replace("**", "\u0000")
    parts = re.split(r"(\u0000[^\u0000]+\u0000|\*[^*]+\*|`[^`]+`)", text)
    for part in parts:
        if not part:
            continue
        run = paragraph.add_run()
        if part.startswith("\u0000") and part.endswith("\u0000"):
            run.text = part[1:-1]
            run.bold = True
        elif part.startswith("*") and part.endswith("*"):
            run.text = part[1:-1]
            run.italic = True
        elif part.startswith("`") and part.endswith("`"):
            run.text = part[1:-1]
            run.font.name = "Consolas"
            run.font.size = Pt(9.5)
        else:
            run.text = part.replace("\u0000", "")


def add_markdown_paragraph(doc: Document, line: str) -> None:
    if not line.strip():
        return
    paragraph = doc.add_paragraph()
    apply_inline_markup(paragraph, line.strip())


def parse_table(lines: list[str]) -> list[list[str]]:
    rows = []
    for line in lines:
        cells = [cell.strip() for cell in line.strip().strip("|").split("|")]
        if all(re.fullmatch(r":?-{3,}:?", cell) for cell in cells):
            continue
        rows.append(cells)
    return rows


def add_table(doc: Document, table_lines: list[str]) -> None:
    rows = parse_table(table_lines)
    if not rows:
        return
    table = doc.add_table(rows=len(rows), cols=max(len(row) for row in rows))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    set_table_width(table)
    for row_idx, row in enumerate(rows):
        for col_idx in range(len(table.columns)):
            cell = table.cell(row_idx, col_idx)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            text = row[col_idx] if col_idx < len(row) else ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if re.fullmatch(r"[\d.%<>=+\- ]+", text) else WD_ALIGN_PARAGRAPH.LEFT
            apply_inline_markup(p, text)
            if row_idx == 0:
                set_cell_shading(cell, "F2F4F7")
                for run in p.runs:
                    run.bold = True
    doc.add_paragraph()


def read_working_title() -> str:
    for line in SOURCE.read_text(encoding="utf-8").splitlines():
        if line.startswith("# "):
            return line[2:].strip()
        if line.startswith("**Working title:**"):
            return re.sub(r"\*\*", "", line.split(":", 1)[1]).strip()
    return "Cytomove: a browser-local and reviewable workflow for scratch wound healing assay quantification"


def add_cover(doc: Document) -> None:
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run(read_working_title())
    run.font.name = "Calibri"
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string("0B2545")

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(12)
    subtitle.add_run(f"Submission manuscript | Generated from {SOURCE.relative_to(ROOT).as_posix()}").italic = True

    meta = doc.add_paragraph()
    meta.add_run("Status: ").bold = True
    meta.add_run("submission build, preliminary validation data\n")
    meta.add_run("Reference file: ").bold = True
    meta.add_run("docs/references/cytomove-preprint.bib\n")
    meta.add_run("Figures: ").bold = True
    meta.add_run("docs/manuscript_figures/")
    doc.add_page_break()


def add_figures(doc: Document, captions: dict[str, str]) -> None:
    doc.add_page_break()
    doc.add_heading("Figures", level=1)
    figure_keys = ["Figure 1", "Figure 2", "Figure 3", "Figure 4", "Figure 5"]
    for index, key in enumerate(figure_keys):
        if index > 0:
            doc.add_page_break()
        path = FIGURE_MAP[key]
        if not path.exists():
            continue
        doc.add_picture(str(optimized_figure(path, max_width_px=1800)), width=Inches(6.4))
        caption = doc.add_paragraph()
        caption.paragraph_format.space_after = Pt(12)
        caption.alignment = WD_ALIGN_PARAGRAPH.LEFT
        apply_inline_markup(caption, captions.get(key, key))

    doc.add_page_break()
    path = FIGURE_MAP["Supplementary Figure S1"]
    if path.exists():
        doc.add_picture(str(optimized_figure(path, max_width_px=1800)), width=Inches(6.4))
        caption = doc.add_paragraph()
        caption.paragraph_format.space_after = Pt(12)
        caption.alignment = WD_ALIGN_PARAGRAPH.LEFT
        apply_inline_markup(caption, captions.get("Supplementary Figure S1", "Supplementary Figure S1"))

    section = doc.add_section()
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    for side in (section.top_margin, section.bottom_margin, section.left_margin, section.right_margin):
        side = side
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.45)
    section.right_margin = Inches(0.45)
    doc.add_heading("Supplementary Landscape Figure", level=1)
    path = FIGURE_MAP["Supplementary Figure S2"]
    if path.exists():
        doc.add_picture(str(optimized_figure(path, max_width_px=3000)), width=Inches(10.0))
        caption = doc.add_paragraph()
        apply_inline_markup(caption, captions.get("Supplementary Figure S2", "Supplementary Figure S2"))


def optimized_figure(path: Path, max_width_px: int) -> Path:
    OPTIMIZED_FIGURES.mkdir(parents=True, exist_ok=True)
    out = OPTIMIZED_FIGURES / f"{path.stem}_{max_width_px}.jpg"
    if out.exists() and out.stat().st_mtime >= path.stat().st_mtime:
        return out
    with Image.open(path) as image:
        image = ImageOps.exif_transpose(image).convert("RGB")
        if image.width > max_width_px:
            ratio = max_width_px / image.width
            next_size = (max_width_px, max(1, int(image.height * ratio)))
            image = image.resize(next_size, Image.Resampling.LANCZOS)
        image.save(out, "JPEG", quality=88, optimize=True, dpi=(220, 220))
    return out


def build_docx() -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    configure_styles(doc)
    add_cover(doc)

    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    in_refs = False
    in_open_todos = False
    captions: dict[str, str] = {}
    table_buffer: list[str] = []
    paragraph_buffer: list[str] = []

    def flush_paragraph() -> None:
        nonlocal paragraph_buffer
        if paragraph_buffer:
            add_markdown_paragraph(doc, " ".join(paragraph_buffer))
            paragraph_buffer = []

    def flush_table() -> None:
        nonlocal table_buffer
        if table_buffer:
            flush_paragraph()
            add_table(doc, table_buffer)
            table_buffer = []

    for raw_line in lines:
        line = raw_line.rstrip()
        if line.startswith("## Figure Captions"):
            flush_table()
            flush_paragraph()
        if line.startswith("## References"):
            flush_table()
            flush_paragraph()
            in_refs = True
        if line.startswith("## Open TODOs"):
            in_refs = False
            in_open_todos = True

        caption_match = re.match(r"\*\*(Figure \d+|Supplementary Figure S\d+)\.([^*]+)\*\*\s*(.*)", line)
        if caption_match:
            label = caption_match.group(1)
            title = caption_match.group(2).strip().rstrip(".")
            body = caption_match.group(3).strip()
            captions[label] = f"{label}. {title}. {body}"

        if line.startswith("# "):
            continue
        if line.startswith("**Working title:**"):
            continue
        if line.startswith("**Status:**") or line.startswith("**Last updated:**") or line.startswith("**Primary data source:**") or line.startswith("**Reference source file:**") or line.startswith("**Preprint positioning:**") or line.startswith("**Positioning:**"):
            continue
        if line.strip() == "---":
            continue

        if line.startswith("|"):
            flush_paragraph()
            table_buffer.append(line)
            continue
        flush_table()

        if not line.strip():
            flush_paragraph()
            continue
        if line.startswith("## "):
            flush_paragraph()
            doc.add_heading(line[3:].strip(), level=1)
            continue
        if line.startswith("### "):
            flush_paragraph()
            doc.add_heading(line[4:].strip(), level=2)
            continue
        if line.startswith("- "):
            flush_paragraph()
            paragraph = doc.add_paragraph(style="List Bullet")
            apply_inline_markup(paragraph, line[2:].strip())
            continue
        if re.match(r"\d+\. ", line):
            flush_paragraph()
            paragraph = doc.add_paragraph(style="List Number")
            apply_inline_markup(paragraph, re.sub(r"^\d+\. ", "", line).strip())
            continue

        paragraph_buffer.append(line)

    flush_table()
    flush_paragraph()
    add_figures(doc, captions)
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build_docx()
